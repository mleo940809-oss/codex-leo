from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_NAME = "Weixin摇一摇_PRD_v2.21_20260513.md"
OUTPUT_NAME = "Weixin摇一摇_PRD_v2.21_20260513.docx"


def find_source() -> Path:
    matches = list(ROOT.rglob(SOURCE_NAME))
    if not matches:
        raise FileNotFoundError(f"Could not find {SOURCE_NAME} under {ROOT}")
    matches.sort(key=lambda p: len(p.parts))
    return matches[0]


def set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_indent(table, indent_dxa: int = 120) -> None:
    tblPr = table._tbl.tblPr
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_in: Sequence[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def shade_header_row(table, fill: str = "E8EEF5") -> None:
    for cell in table.rows[0].cells:
        set_cell_shading(cell, fill)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color="000000") -> None:
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, target: str, color: str = "0563C1", underline: bool = True):
    part = paragraph.part
    target = target.strip().strip('"')
    if re.match(r"^[A-Za-z]:[\\/]", target):
        try:
            target = Path(target).resolve().as_uri()
        except Exception:
            target = target.replace("\\", "/")
            if not target.startswith("file:///"):
                target = "file:///" + target.lstrip("/")

    r_id = part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "1"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(1))
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal_rPr = normal._element.get_or_add_rPr()
    normal_rFonts = normal_rPr.rFonts
    if normal_rFonts is None:
        normal_rFonts = OxmlElement("w:rFonts")
        normal_rPr.append(normal_rFonts)
    normal_rFonts.set(qn("w:ascii"), "Calibri")
    normal_rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        s_rPr = style._element.get_or_add_rPr()
        s_rFonts = s_rPr.rFonts
        if s_rFonts is None:
            s_rFonts = OxmlElement("w:rFonts")
            s_rPr.append(s_rFonts)
        s_rFonts.set(qn("w:ascii"), "Calibri")
        s_rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name, size, before, after in [
        ("Heading 4", 11, 8, 4),
        ("Heading 5", 10.5, 6, 3),
        ("Heading 6", 10, 4, 2),
    ]:
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, 1)
        style.font.name = "Calibri"
        s_rPr = style._element.get_or_add_rPr()
        s_rFonts = s_rPr.rFonts
        if s_rFonts is None:
            s_rFonts = OxmlElement("w:rFonts")
            s_rPr.append(s_rFonts)
        s_rFonts.set(qn("w:ascii"), "Calibri")
        s_rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("1F4D78")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Number"):
        if style_name in styles:
            s = styles[style_name]
            s.font.name = "Calibri"
            s_rPr = s._element.get_or_add_rPr()
            s_rFonts = s_rPr.rFonts
            if s_rFonts is None:
                s_rFonts = OxmlElement("w:rFonts")
                s_rPr.append(s_rFonts)
            s_rFonts.set(qn("w:ascii"), "Calibri")
            s_rFonts.set(qn("w:hAnsi"), "Calibri")
            s.font.size = Pt(11)
            s.paragraph_format.space_after = Pt(4)
            s.paragraph_format.line_spacing = 1.25


def add_footer(doc: Document) -> None:
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    run = p.add_run("Weixin摇一摇 PRD  |  Page ")
    set_run_font(run, size=9, color="666666")
    add_page_number(p)


def strip_md(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[(.*?)\]\((.*?)\)", r"\1", text)
    return text.strip()


def split_table_row(line: str) -> List[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells = []
    buf = []
    depth = 0
    for ch in line:
        if ch == "[":
            depth += 1
        elif ch == "]" and depth:
            depth -= 1
        if ch == "|" and depth == 0:
            cells.append("".join(buf).strip())
            buf = []
        else:
            buf.append(ch)
    cells.append("".join(buf).strip())
    return cells


def is_table_sep(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^\|?(\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?$", s))


def guess_table_widths(rows: Sequence[Sequence[str]]) -> List[float]:
    ncols = max(len(r) for r in rows)
    scores = [0.0] * ncols
    for row in rows:
        for i in range(ncols):
            cell = strip_md(row[i]) if i < len(row) else ""
            score = max(len(cell), len(cell.encode("utf-8")) / 3.2)
            scores[i] = max(scores[i], score)
    total = sum(scores) or ncols
    usable = 6.5
    min_w = 0.55
    widths = [max(min_w, usable * (s / total)) for s in scores]
    overflow = sum(widths) - usable
    if overflow > 0:
        adjustable = sum(max(w - min_w, 0) for w in widths)
        if adjustable > 0:
            widths = [w - overflow * max(w - min_w, 0) / adjustable for w in widths]
    return widths


def add_inline_runs(paragraph, text: str, font_name="Calibri", font_size=11, default_color="000000") -> None:
    i = 0
    while i < len(text):
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j != -1:
                run = paragraph.add_run(text[i + 1 : j])
                set_run_font(run, name="Consolas", size=font_size, color="444444")
                i = j + 1
                continue
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j != -1:
                run = paragraph.add_run(text[i + 2 : j])
                set_run_font(run, name=font_name, size=font_size, bold=True, color=default_color)
                i = j + 2
                continue
        if text[i] == "*" and not text.startswith("**", i):
            j = text.find("*", i + 1)
            if j != -1:
                run = paragraph.add_run(text[i + 1 : j])
                set_run_font(run, name=font_name, size=font_size, italic=True, color=default_color)
                i = j + 1
                continue
        if text[i] == "[":
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", text[i:])
            if m:
                label, target = m.group(1), m.group(2)
                add_hyperlink(paragraph, label, target)
                i += m.end()
                continue
        j = i + 1
        while j < len(text) and text[j] not in "`*[":
            j += 1
        run = paragraph.add_run(text[i:j])
        set_run_font(run, name=font_name, size=font_size, color=default_color)
        i = j


def add_text_paragraph(doc: Document, text: str, style: Optional[str] = None, *,
                       before=0, after=6, line=1.25, align=WD_ALIGN_PARAGRAPH.LEFT,
                       font_size=11, bold=False, italic=False, color="000000") -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if text:
        add_inline_runs(p, text, font_size=font_size, default_color=color)
        if bold or italic:
            for run in p.runs:
                run.bold = bold or run.bold
                run.italic = italic or run.italic


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=3, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=22, bold=True, color="000000")


def add_quote_block(doc: Document, lines: Sequence[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        run = p.add_run(line.strip().lstrip("> ").rstrip())
        set_run_font(run, size=9.5, italic=True, color="666666")


def add_list_paragraph(doc: Document, text: str, numbered: bool, level: int = 0) -> None:
    style_name = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style_name)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25 + 0.25 * level)
    pf.first_line_indent = Inches(-0.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    add_inline_runs(p, text)


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        run = p.add_run(line.rstrip())
        set_run_font(run, name="Consolas", size=10, color="444444")


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    set_table_indent(table, 120)
    widths = guess_table_widths(rows)
    set_table_widths(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx in range(ncols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            if c_idx < len(row):
                cell_text = row[c_idx].strip()
                # rebuild cell content with simple inline formatting and local links
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, before=0, after=0, line=1.15)
                if cell_text:
                    raw_path = cell_text.strip().strip('"')
                    if re.match(r"^[A-Za-z]:[\\/].+\.(?:png|jpg|jpeg|gif|drawio|md|pdf)$", raw_path, re.I):
                        label = Path(raw_path).name
                        add_hyperlink(p, label, raw_path)
                    else:
                        add_inline_runs(p, cell_text, font_size=10)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    shade_header_row(table, "E8EEF5")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc: Document, image_path: str, caption: Optional[str] = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=3, line=1.0)
    target = Path(image_path.strip().strip('"'))
    if re.match(r"^[A-Za-z]:[\\/]", str(target)):
        path = target
    else:
        path = (ROOT / target).resolve()
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(cp, before=0, after=4, line=1.0)
            run = cp.add_run(caption)
            set_run_font(run, size=9, italic=True, color="666666")
    else:
        add_text_paragraph(doc, f"[Missing image: {path.name}]", before=0, after=4, color="9B1C1C")


def paragraph_is_image(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^!\[[^\]]*\]\([^)]+\)$", s))


def extract_image_path(line: str) -> str:
    m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
    return m.group(2) if m else ""


def add_mermaid_block(doc: Document, lines: Sequence[str]) -> None:
    boxed = doc.add_table(rows=1, cols=1)
    boxed.style = "Table Grid"
    set_table_indent(boxed, 120)
    set_table_widths(boxed, [6.5])
    cell = boxed.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=0, line=1.05)
    for line in lines:
        run = p.add_run(line.rstrip() + "\n")
        set_run_font(run, name="Consolas", size=9, color="444444")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_and_build(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    saw_title = False
    pending_quote: List[str] = []

    def flush_quote():
        nonlocal pending_quote
        if pending_quote:
            add_quote_block(doc, pending_quote)
            pending_quote = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not saw_title and stripped.startswith("# "):
            add_title(doc, stripped[2:].strip())
            saw_title = True
            i += 1
            continue

        if stripped.startswith(">"):
            pending_quote.append(line)
            i += 1
            continue
        else:
            flush_quote()

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("```"):
            fence = stripped
            block = [line]
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            if i < len(lines):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level <= 3:
                p = doc.add_paragraph(style=f"Heading {level}")
            else:
                p = doc.add_paragraph(style=f"Heading {level}")
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
            add_inline_runs(p, text, font_size={4: 11, 5: 10.5, 6: 10}.get(level, 11), default_color="000000")
            i += 1
            continue

        if paragraph_is_image(stripped):
            image_path = extract_image_path(stripped)
            add_image(doc, image_path)
            i += 1
            continue

        if i + 1 < len(lines) and is_table_sep(lines[i + 1]) and "|" in line:
            rows: List[List[str]] = [split_table_row(line)]
            i += 2
            while i < len(lines):
                nxt = lines[i]
                if not nxt.strip():
                    break
                if "|" not in nxt:
                    break
                if re.match(r"^\s*[-*+]\s+", nxt.strip()):
                    break
                rows.append(split_table_row(nxt))
                i += 1
            add_table(doc, rows)
            continue

        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            item = list_match.group(3).rstrip()
            numbered = marker.endswith(".")
            level = max(indent // 2, 0)
            add_list_paragraph(doc, item, numbered=numbered, level=level)
            i += 1
            continue

        # paragraph merge until the next structural line
        block = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt_s = nxt.strip()
            if not nxt_s:
                break
            if nxt_s.startswith((">", "#", "```")) or paragraph_is_image(nxt_s):
                break
            if i + 1 < len(lines) and is_table_sep(lines[i + 1]) and "|" in nxt:
                break
            if re.match(r"^(\s*)([-*+]|\d+\.)\s+", nxt):
                break
            block.append(nxt_s)
            i += 1
        para_text = " ".join(block)
        add_text_paragraph(doc, para_text)

    flush_quote()


def configure_footer(doc: Document) -> None:
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    run = p.add_run("Weixin摇一摇 PRD  |  Page ")
    set_run_font(run, size=9, color="666666")
    add_page_number(p)


def main() -> None:
    source = find_source()
    output = source.with_name(OUTPUT_NAME)
    text = source.read_text(encoding="utf-8-sig")

    doc = Document()
    configure_document(doc)
    configure_footer(doc)
    parse_and_build(doc, text)
    doc.core_properties.title = "Weixin摇一摇 PRD"
    doc.core_properties.subject = "Markdown to Word conversion"
    doc.core_properties.author = "Codex"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
